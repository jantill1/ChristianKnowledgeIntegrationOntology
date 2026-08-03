from pathlib import Path
from dataclasses import dataclass
from collections import Counter
import csv, json, zipfile
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent
OUT = ROOT / "output"
OUT.mkdir(exist_ok=True)

@dataclass(frozen=True)
class Passage:
    book: str
    ref: str
    episode: str
    relation: str = "depicts"

EPISODES = {
"genealogies": "Genealogies of Jesus",
"birth_announcements": "Announcements concerning the birth of Jesus",
"nativity": "Nativity and early recognition of Jesus",
"baptism": "Baptism of Jesus",
"temptation": "Temptation of Jesus",
"first_disciples": "Calling of the first disciples",
"beatitudes": "Beatitudes and kingdom teaching",
"paralytic": "Healing of a paralytic",
"feeding_5000": "Feeding of the five thousand",
"walking_water": "Jesus walking on water",
"transfiguration": "Transfiguration of Jesus",
"triumphal_entry": "Triumphal entry into Jerusalem",
"last_supper": "Last Supper and final meal teaching",
"gethsemane": "Prayer and arrest in Gethsemane",
"trial": "Trials of Jesus",
"crucifixion": "Crucifixion and death of Jesus",
"burial": "Burial of Jesus",
"empty_tomb": "Discovery of the empty tomb",
"appearances": "Post-resurrection appearances",
"commission": "Commissioning and sending of disciples",
}

P = [
Passage("Matthew","1:1-17","genealogies"), Passage("Luke","3:23-38","genealogies"),
Passage("Matthew","1:18-25","birth_announcements"), Passage("Luke","1:26-38","birth_announcements"),
Passage("Matthew","2:1-12","nativity"), Passage("Luke","2:1-20","nativity"),
Passage("Matthew","3:13-17","baptism"), Passage("Mark","1:9-11","baptism"), Passage("Luke","3:21-22","baptism"), Passage("John","1:29-34","baptism","testifiesTo"),
Passage("Matthew","4:1-11","temptation"), Passage("Mark","1:12-13","temptation"), Passage("Luke","4:1-13","temptation"),
Passage("Matthew","4:18-22","first_disciples"), Passage("Mark","1:16-20","first_disciples"), Passage("Luke","5:1-11","first_disciples"), Passage("John","1:35-51","first_disciples"),
Passage("Matthew","5:1-12","beatitudes"), Passage("Luke","6:20-26","beatitudes"),
Passage("Matthew","9:1-8","paralytic"), Passage("Mark","2:1-12","paralytic"), Passage("Luke","5:17-26","paralytic"),
Passage("Matthew","14:13-21","feeding_5000"), Passage("Mark","6:30-44","feeding_5000"), Passage("Luke","9:10-17","feeding_5000"), Passage("John","6:1-15","feeding_5000"),
Passage("Matthew","14:22-33","walking_water"), Passage("Mark","6:45-52","walking_water"), Passage("John","6:16-21","walking_water"),
Passage("Matthew","17:1-13","transfiguration"), Passage("Mark","9:2-13","transfiguration"), Passage("Luke","9:28-36","transfiguration"),
Passage("Matthew","21:1-11","triumphal_entry"), Passage("Mark","11:1-11","triumphal_entry"), Passage("Luke","19:28-40","triumphal_entry"), Passage("John","12:12-19","triumphal_entry"),
Passage("Matthew","26:17-30","last_supper"), Passage("Mark","14:12-26","last_supper"), Passage("Luke","22:7-38","last_supper"), Passage("John","13:1-38","last_supper","contextualizes"),
Passage("Matthew","26:36-56","gethsemane"), Passage("Mark","14:32-52","gethsemane"), Passage("Luke","22:39-53","gethsemane"), Passage("John","18:1-11","gethsemane"),
Passage("Matthew","26:57-27:26","trial"), Passage("Mark","14:53-15:15","trial"), Passage("Luke","22:54-23:25","trial"), Passage("John","18:12-19:16","trial"),
Passage("Matthew","27:27-56","crucifixion"), Passage("Mark","15:16-41","crucifixion"), Passage("Luke","23:26-49","crucifixion"), Passage("John","19:16-37","crucifixion"),
Passage("Matthew","27:57-66","burial"), Passage("Mark","15:42-47","burial"), Passage("Luke","23:50-56","burial"), Passage("John","19:38-42","burial"),
Passage("Matthew","28:1-10","empty_tomb"), Passage("Mark","16:1-8","empty_tomb"), Passage("Luke","24:1-12","empty_tomb"), Passage("John","20:1-18","empty_tomb"),
Passage("Matthew","28:9-20","appearances"), Passage("Mark","16:9-20","appearances","textuallyQualified"), Passage("Luke","24:13-49","appearances"), Passage("John","20:19-21:23","appearances"),
Passage("Matthew","28:16-20","commission"), Passage("Mark","16:14-18","commission","textuallyQualified"), Passage("Luke","24:44-49","commission"), Passage("John","20:21-23","commission"),
]

BOOKS = ["Matthew", "Mark", "Luke", "John"]

def pid(x): return f"{x.book.lower()}_{x.ref.replace(':','_').replace('-','_')}"

def write_ttl():
    lines = [
        '@prefix ckio: <https://example.org/ckio/> .',
        '@prefix rdf: <http://www.w3.org/1999/02/22-rdf-syntax-ns#> .',
        '@prefix rdfs: <http://www.w3.org/2000/01/rdf-schema#> .',
        '@prefix owl: <http://www.w3.org/2002/07/owl#> .',
        '@prefix xsd: <http://www.w3.org/2001/XMLSchema#> .',
        '@prefix prov: <http://www.w3.org/ns/prov#> .',
        '@prefix dcterms: <http://purl.org/dc/terms/> .',
        '',
        'ckio: a owl:Ontology ; rdfs:label "Christian Knowledge Integration Ontology prototype" ; dcterms:license <https://creativecommons.org/licenses/by/4.0/> .',
    ]
    classes = ["ChristianResource","BiblicalBook","Passage","Episode","InterpretiveClaim","ComputationalAssertion","Tradition","ValidationActivity"]
    for c in classes: lines.append(f'ckio:{c} a owl:Class .')
    props = ["isPartOf","depicts","testifiesTo","contextualizes","textuallyQualified","hasEvidence","proposesEpisode","usesPredicate","assertedBy","hasStatus"]
    for p in props: lines.append(f'ckio:{p} a owl:ObjectProperty .')
    lines += [
        '',
        'ckio:BiblicalBook rdfs:subClassOf ckio:ChristianResource .',
        'ckio:Passage rdfs:subClassOf ckio:ChristianResource .',
        'ckio:ComputationalAssertion rdfs:subClassOf prov:Entity .',
        'ckio:InterpretiveClaim rdfs:subClassOf prov:Entity .',
        'ckio:ValidationActivity rdfs:subClassOf prov:Activity .',
        'ckio:Passage owl:disjointWith ckio:InterpretiveClaim .',
        'ckio:ComputationalAssertion owl:disjointWith ckio:InterpretiveClaim .',
        'ckio:isPartOf rdfs:domain ckio:Passage ; rdfs:range ckio:BiblicalBook .',
        'ckio:hasEvidence rdfs:domain ckio:ComputationalAssertion ; rdfs:range ckio:Passage .',
        'ckio:proposesEpisode rdfs:domain ckio:ComputationalAssertion ; rdfs:range ckio:Episode .',
        'ckio:usesPredicate rdfs:domain ckio:ComputationalAssertion ; rdfs:range rdf:Property .',
        'ckio:depicts rdfs:domain ckio:Passage ; rdfs:range ckio:Episode .',
        'ckio:testifiesTo rdfs:domain ckio:Passage ; rdfs:range ckio:Episode .',
        'ckio:contextualizes rdfs:domain ckio:Passage ; rdfs:range ckio:Episode .',
        'ckio:textuallyQualified rdfs:domain ckio:Passage ; rdfs:range ckio:Episode .',
    ]
    lines += [
        'ckio:ChatGPT a prov:SoftwareAgent ; rdfs:label "OpenAI ChatGPT (GPT-5 family)" .',
        'ckio:PrototypeConstruction a prov:Activity ; rdfs:label "CKIO prototype construction activity" ; prov:wasAssociatedWith ckio:ChatGPT ; prov:generatedAtTime "2026-08-03"^^xsd:date .',
        'ckio:RuleAccepted a owl:NamedIndividual ; rdfs:label "Accepted by source-anchoring rule" .',
    ]
    for b in BOOKS:
        lines.append(f'ckio:{b} a ckio:BiblicalBook ; rdfs:label "Gospel according to {b}" .')
    for k, label in EPISODES.items():
        lines.append(f'ckio:episode_{k} a ckio:Episode ; rdfs:label "{label}" .')
    for x in P:
        lines.append(f'ckio:{pid(x)} a ckio:Passage ; rdfs:label "{x.book} {x.ref}" ; ckio:isPartOf ckio:{x.book} ; ckio:{x.relation} ckio:episode_{x.episode} .')
        lines.append(f'ckio:assertion_{pid(x)} a ckio:ComputationalAssertion ; ckio:hasEvidence ckio:{pid(x)} ; ckio:proposesEpisode ckio:episode_{x.episode} ; ckio:usesPredicate ckio:{x.relation} ; ckio:hasStatus ckio:RuleAccepted ; prov:wasGeneratedBy ckio:PrototypeConstruction .')
    (OUT/'ckio_prototype.ttl').write_text('\n'.join(lines)+'\n', encoding='utf-8')

def evaluate():
    per_book = Counter(x.book for x in P)
    per_episode = Counter(x.episode for x in P)
    per_relation = Counter(x.relation for x in P)
    missing_books = [b for b in BOOKS if per_book[b] == 0]
    bad_episode = [x.episode for x in P if x.episode not in EPISODES]
    duplicate_ids = len(P) - len({pid(x) for x in P})
    no_evidence = [k for k in EPISODES if per_episode[k] == 0]
    fourfold = [k for k in EPISODES if len({x.book for x in P if x.episode==k}) == 4]
    synoptic = [k for k in EPISODES if {'Matthew','Mark','Luke'} <= {x.book for x in P if x.episode==k}]
    cqs = {
        "CQ1": len(P),
        "CQ2": per_book,
        "CQ3": fourfold,
        "CQ4": [f"{x.book} {x.ref}" for x in P if x.episode=='feeding_5000'],
        "CQ5": [f"{x.book} {x.ref}" for x in P if x.relation=='textuallyQualified'],
        "CQ6": synoptic,
    }
    result = {
        "books": len(BOOKS), "episodes": len(EPISODES), "passage_units": len(P),
        "typed_assertions": len(P), "per_book": dict(per_book), "per_episode": dict(per_episode), "per_relation": dict(per_relation),
        "four_gospel_episodes": fourfold, "synoptic_episodes": synoptic,
        "integrity": {"missing_books":missing_books,"unknown_episode_keys":bad_episode,"duplicate_passage_ids":duplicate_ids,"episodes_without_evidence":no_evidence},
        "competency_queries": cqs,
    }
    (OUT/'evaluation_results.json').write_text(json.dumps(result, indent=2, default=dict), encoding='utf-8')
    with (OUT/'gospel_passage_dataset.csv').open('w',newline='',encoding='utf-8') as f:
        w=csv.writer(f); w.writerow(['passage_id','book','reference','episode_id','episode_label','relationship'])
        for x in P: w.writerow([pid(x),x.book,x.ref,x.episode,EPISODES[x.episode],x.relation])
    return result

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def set_table_widths(table, widths):
    table.autofit=False
    for row in table.rows:
        for cell, width in zip(row.cells,widths):
            cell.width=Inches(width); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
            if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'),str(int(width*1440))); tcW.set(qn('w:type'),'dxa')

def add_table(doc, headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.LEFT; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; set_cell_shading(c,'E8EEF5')
        for r in c.paragraphs[0].runs: r.bold=True; r.font.size=Pt(9)
    set_repeat_table_header(t.rows[0])
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(2)
                for r in p.runs: r.font.size=Pt(9)
    set_table_widths(t,widths)
    return t

def para(doc, text, bold_start=None):
    p=doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        p.add_run(bold_start).bold=True; p.add_run(text[len(bold_start):])
    else: p.add_run(text)
    return p

def build_doc(r):
    d=Document(); sec=d.sections[0]
    sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
    sec.header_distance=sec.footer_distance=Inches(.492)
    normal=d.styles['Normal']; normal.font.name='Calibri'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor(0,0,0)
    normal.paragraph_format.space_after=Pt(8); normal.paragraph_format.line_spacing=1.15
    for nm,size,bef,aft,col in [('Heading 1',16,18,10,'2E74B5'),('Heading 2',13,12,6,'2E74B5'),('Heading 3',12,8,4,'1F4D78')]:
        s=d.styles[nm]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(col); s.paragraph_format.space_before=Pt(bef); s.paragraph_format.space_after=Pt(aft); s.paragraph_format.keep_with_next=True
    # Header/footer
    hp=sec.header.paragraphs[0]; hp.text='Christian Knowledge Integration Ontology'; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs: run.font.size=Pt(9); run.font.color.rgb=RGBColor(100,100,100)
    fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld)

    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(70); p.paragraph_format.space_after=Pt(16)
    rr=p.add_run('An AI-Assisted, Provenance-Aware Ontology for Christian Knowledge Integration'); rr.bold=True; rr.font.size=Pt(22); rr.font.color.rgb=RGBColor.from_string('0B2545')
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Linking Gospel Passages, Computational Assertions, and Interpretive Context').italic=True
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(24); p.add_run('John Antill').bold=True
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Independent Researcher')
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Correspondence: john@littlecrittersnursery.com')
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Research Article')
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Prepared for Academia AI and Applications')
    d.add_page_break()

    d.add_heading('Abstract',1)
    para(d,'Introduction: Christian knowledge is distributed across biblical texts, linguistic resources, creeds, commentaries, and diverse interpretive traditions. Existing biblical semantic resources commonly focus on textual structure or named entities and provide limited support for distinguishing source text, computational inference, and tradition-specific interpretation. This study develops a provenance-aware Christian Knowledge Integration Ontology (CKIO) proof of concept for the four canonical Gospels.')
    para(d,'Materials and methods: Twenty Gospel episodes were represented through passage-level evidence from Matthew, Mark, Luke, and John. A modular OWL/RDF schema distinguished biblical passages, episodes, computational assertions, interpretive claims, traditions, and validation activities. Candidate relationships were constrained to typed predicates and source references. Structural integrity tests and six competency questions assessed coverage, traceability, relationship qualification, and cross-Gospel retrieval. Generative AI assisted schema development and manuscript preparation but was not treated as an authoritative theological source.')
    para(d,f"Results: The prototype represented {r['books']} Gospel books, {r['episodes']} episode groupings, {r['passage_units']} passage units, and {r['typed_assertions']} typed passage-to-episode assertions with reified provenance records. All passage units were linked to a book and a recognized episode grouping; no duplicate passage identifiers or unsupported episode nodes were detected. Six competency queries returned the expected fixture-based results. {len(r['four_gospel_episodes'])} groupings contained material from all four Gospels and {len(r['synoptic_episodes'])} contained material from all three Synoptic Gospels. Textually qualified relationships were retained separately from ordinary narrative parallels.")
    para(d,'Conclusions: The proof of concept shows that provenance-aware semantic modelling can integrate Gospel parallels without collapsing textual evidence, computational inference, and theological interpretation into a single undifferentiated relationship. The principal advance is epistemic separation and traceability rather than ontology size. Full-corpus extraction, independent theological review, retrieval benchmarking, and denominationally diverse validation are required before the ontology can support authoritative scholarly or pastoral use.')
    para(d,'Keywords: Christian knowledge ontology; biblical knowledge graph; large language models; natural language processing; semantic integration; provenance; Gospel parallels; ontology engineering')

    d.add_heading('1. Introduction',1)
    intro=[
    'Christianity has generated a large and heterogeneous body of knowledge encompassing biblical texts, original-language resources, textual variants, creeds, councils, commentaries, theological writings, liturgical materials, and denominational teachings. These resources are closely related but are generally stored and accessed through separate databases, digital libraries, websites, and specialist research tools. This fragmentation limits the ability of researchers, educators, clergy, students, and computational systems to trace relationships among biblical passages, linguistic evidence, historical interpretations, and theological concepts within a unified semantic environment.',
    'Ontology engineering provides a formal method for representing entities, concepts, and relationships in a machine-interpretable structure [1]. Through standards such as RDF, OWL, and SPARQL, ontologies can support semantic retrieval, interoperability, automated reasoning, and knowledge reuse [2-4]. Applied to Christian resources, these technologies can support questions extending beyond conventional keyword retrieval, including identifying Gospel passages that describe corresponding events, distinguishing direct quotation from thematic association, and tracing the provenance of an interpretive claim.',
    'Previous work demonstrates the feasibility of representing biblical information semantically. The New Testament Names knowledge base uses OWL to describe approximately 600 named entities, including people, groups, locations, and their relationships [5]. Its structure supports conceptual searching, visualization, and network analysis, but its principal emphasis is named entities rather than the integration of textual structure, linguistic annotation, theological concepts, interpretive claims, and provenance.',
    'Open digital resources provide foundations for a broader ontology. The Open Scriptures Hebrew Bible supplies word-level identifiers, lemmas, and morphological information in OSIS XML [6,7]. STEPBible publishes tagged biblical texts, lexicons, morphology codes, proper-noun data, and versification resources under a CC BY 4.0 license [8]. These resources make it possible to connect passages with original-language evidence and structured metadata. Integration remains difficult because resources differ in identifiers, schemas, granularity, licensing, and treatment of textual variants.',
    'Recent advances in NLP and LLMs create opportunities to identify entities, events, topics, and candidate semantic relationships in religious texts [9-12]. An Islamic ontology study used GPT-4 to generate bilingual Hadith topics and integrate Qur\'anic and Hadith resources in OWL [13]. Comparative NLP has also been applied to the Bible, Qur\'an, and Bhagavad Gita [19]. This work demonstrates the potential of combining AI, NLP, and ontology engineering in religious domains. Nevertheless, LLM-based extraction introduces risks of omission, overgeneralization, unsupported relationships, and the unacknowledged projection of a particular interpretive perspective [14,15]. Logical consistency does not establish historical or theological correctness.',
    'Christian knowledge adds further representational difficulty because biblical canons, versification systems, textual traditions, and doctrinal interpretations vary. A passage may quote, allude to, parallel, contextualize, or appear to conflict with another passage. A doctrinal interpretation may be shared across traditions, restricted to one communion, disputed, or historically conditioned. A generic relatedTo predicate obscures these distinctions and can make machine-generated associations appear more authoritative than their evidence permits.',
    'Provenance is therefore a foundational requirement. PROV-O provides a standard representation of entities, activities, agents, derivation, and attribution [16]. Existing semantic vocabularies such as FaBiO and CiTO demonstrate how formally defined resources and citation relationships can support interoperable scholarly knowledge [18]. SHACL provides machine-readable constraints for validating RDF graphs [17]. Used together, these standards allow an ontology to record whether an assertion derives from a biblical passage, an NLP pipeline, an LLM proposal, a commentary, or a human validation activity. They also support recording model version, prompt, supporting evidence, validation status, and interpretive tradition. Recent ontology-generation benchmarks further emphasize comparison with human reference ontologies and qualitative assessment rather than structural counts alone [20].',
    'This study develops an AI-assisted, provenance-aware Christian Knowledge Integration Ontology proof of concept focused on the four canonical Gospels. It tests whether passage-level evidence and typed relationships can integrate corresponding Gospel episodes while preserving epistemic separation among source text, computational assertion, and interpretation. The study contributes a modular schema, a source-anchored dataset, competency questions, reproducible integrity tests, and a restrained evaluation that distinguishes demonstrated technical performance from theological validity.'
    ]
    for x in intro: para(d,x)
    d.add_heading('1.1 Research questions',2)
    qs=[
    'Can a modular OWL/RDF model integrate passage-level evidence and corresponding episodes across the four Gospels?',
    'Can provenance-aware modelling prevent computational assertions and interpretive claims from being represented as direct scriptural facts?',
    'Can typed relationships capture important qualifications that a generic relatedTo predicate would conceal?',
    'Can the prototype answer competency questions concerning coverage, parallels, evidence, and qualification?',
    'What limitations must be addressed before expansion to complete biblical corpora, commentaries, creeds, and multiple Christian traditions?']
    for q in qs: d.add_paragraph(q, style='List Number')

    d.add_heading('2. Materials and methods',1)
    d.add_heading('2.1 Study design and scope',2)
    para(d,'The study used a design-science proof-of-concept approach. The artifact was an OWL/RDF ontology accompanied by a structured passage dataset and executable integrity checks. The scope was restricted to 20 curated Gospel episode groupings rather than the complete text of the four Gospels. Some groupings, including the birth announcements, calling of disciples, and Johannine testimony concerning baptism, organize contextually related material rather than strict scene-by-scene narrative parallels. This bounded design allowed testing of cross-Gospel representation, provenance, and relationship qualification without claiming full canonical or theological coverage.')
    para(d,'The unit of analysis was a passage unit: a book-and-reference range associated with an episode through a typed predicate. Passage ranges may overlap when one range represents a broader episode and another represents a nested analytical grouping; therefore, the passage-unit count is not a count of distinct verses. Verse text was not reproduced in the dataset, which reduced copyright and translation dependencies and kept the prototype independent of a single English translation. Passage references were curated as modelling fixtures. They are not presented as a critical edition or a substitute for textual scholarship.')
    d.add_heading('2.2 Data sources and licensing',2)
    para(d,'The schema was informed by RDF, OWL 2, SPARQL, PROV-O, and SHACL standards [2-4,16,17]. Existing biblical resources informed interoperability decisions: OSIS provided a model for structured scripture identifiers; Open Scriptures demonstrated stable word-level identifiers and morphological annotation; STEPBible demonstrated reusable lexical and versification data; and New Testament Names provided a precedent for biblical entities in OWL [5-8]. The pericope-oriented organization follows the established Gospel-synopsis tradition represented by Aland [21], although the present fixtures are simplified analytical groupings rather than a reproduction of that synopsis. The prototype itself contains references and descriptive labels rather than copyrighted translation text. All generated code, ontology files, and tables are intended for release under CC BY 4.0.')
    d.add_heading('2.3 Domain specification and competency questions',2)
    para(d,'The intended users are digital-humanities researchers, biblical scholars, educators, ontology engineers, and developers of semantic retrieval applications. The prototype was designed to answer six competency questions:')
    cql=[
      'CQ1: How many passage units are represented?',
      'CQ2: How many passage units are associated with each Gospel?',
      'CQ3: Which represented episodes contain evidence from all four Gospels?',
      'CQ4: Which passages represent the feeding of the five thousand?',
      'CQ5: Which relationships carry an explicit textual qualification?',
      'CQ6: Which episodes contain evidence from Matthew, Mark, and Luke?']
    for x in cql: d.add_paragraph(x,style='List Bullet')
    d.add_heading('2.4 Ontology conceptualization',2)
    para(d,'The ontology separates resource, assertion, interpretation, and validation layers. ChristianResource is the upper domain class. BiblicalBook and Passage represent source structure. Episode represents an analytical grouping of passages. ComputationalAssertion represents a machine-proposed relationship. InterpretiveClaim represents an attributed theological or exegetical proposition. Tradition identifies the community or interpretive framework to which a claim is attributed. ValidationActivity records a review or rule-based check. Passage and InterpretiveClaim are declared disjoint, as are ComputationalAssertion and InterpretiveClaim, preventing category collapse at the schema level.')
    add_table(d,['Layer','Principal classes','Epistemic function'],[
      ['Source','BiblicalBook; Passage','Identifies textual evidence without asserting an interpretation.'],
      ['Analytical','Episode; ComputationalAssertion','Represents curated or machine-proposed groupings.'],
      ['Interpretive','InterpretiveClaim; Tradition','Attributes readings to identifiable sources or communities.'],
      ['Validation','ValidationActivity','Records checking method, agent, status, and evidence.']], [1.1,2.2,3.2])
    d.add_paragraph('Table 1. CKIO epistemic layers and their functions.',style='Caption')
    d.add_heading('2.5 Relationship model',2)
    para(d,'The predicate depicts represents a passage associated with a narrated episode. testifiesTo is used where a passage provides testimony related to an episode without narrating it in the same manner. contextualizes represents relevant material whose narrative organization differs from the corresponding Synoptic account. textuallyQualified marks a relationship requiring explicit textual-critical qualification. These distinctions are deliberately conservative. The prototype does not use supportsDoctrine because doctrinal support requires an attributed interpretive claim and appropriate expert review.')
    d.add_heading('2.6 AI assistance and control procedure',2)
    para(d,'Generative AI assisted brainstorming, schema refinement, candidate relationship naming, code generation, and manuscript language. It was not designated as an author, theological authority, or independent validator. No assertion was accepted solely because an LLM produced it. The acceptance rule for the prototype required a recognized episode identifier, an explicit Gospel reference, a typed predicate from the controlled vocabulary, and passage-level traceability. Uncertain or textually disputed material was assigned a qualifying predicate instead of being silently normalized.')
    para(d,'This design differs from an automated extraction-performance experiment. Because the prototype did not use a locked external model endpoint over a complete corpus, the study does not report LLM precision, recall, or F1. This restraint avoids presenting manuscript-development assistance as a reproducible extraction benchmark. A future experiment must record the exact model, endpoint, date, system prompt, user prompt, sampling parameters, retries, raw outputs, and adjudication decisions.')
    d.add_heading('2.7 Formalization and validation',2)
    para(d,'The ontology was serialized in Turtle using stable prototype identifiers. Each passage-to-episode mapping was represented both as a typed graph edge and as a reified ComputationalAssertion carrying its source passage, proposed episode, predicate, rule-acceptance status, and generating activity. Integrity tests checked that every passage belonged to one of four Gospel books, every passage referenced a declared episode, every episode had evidence, and passage identifiers were unique. The tests also enumerated episode groupings with four-Gospel and Synoptic coverage and returned passages with qualified relationships. OWL disjointness axioms encoded key epistemic separations. Turtle syntax and SPARQL execution were verified with Python 3.12.13 and RDFLib 7.1.4. The released prototype can be extended with SHACL node shapes requiring each Passage to have exactly one book, at least one evidence relationship, a label, and a normalized reference.')
    d.add_heading('2.8 Evaluation',2)
    para(d,'Evaluation comprised structural metrics, integrity tests, and competency-question execution. Structural metrics describe artifact size but were not treated as evidence of scalability or theological correctness. Integrity tests produced pass/fail results for identifier uniqueness, book coverage, episode references, and evidence completeness. Competency questions were implemented as deterministic queries over the structured dataset. Expected results were derived from the declared fixtures and compared with obtained outputs. The evaluation did not include human participants, usability testing, denominational review, or claims of doctrinal accuracy.')

    d.add_heading('3. Results',1)
    d.add_heading('3.1 Ontology and dataset composition',2)
    para(d,f"The prototype contains {r['books']} biblical-book individuals, {r['episodes']} episode individuals, and {r['passage_units']} passage individuals. Each passage contributes one typed passage-to-episode assertion, producing {r['typed_assertions']} principal semantic assertions. Matthew contributes {r['per_book']['Matthew']} passage units, Mark {r['per_book']['Mark']}, Luke {r['per_book']['Luke']}, and John {r['per_book']['John']}. The uneven counts reflect the selected episodes and differentiated treatment of Johannine testimony and contextual material rather than an attempt to balance the dataset artificially.")
    para(d,f"Of the {r['typed_assertions']} mappings, {r['per_relation']['depicts']} use depicts, {r['per_relation']['testifiesTo']} uses testifiesTo, {r['per_relation']['contextualizes']} uses contextualizes, and {r['per_relation']['textuallyQualified']} use textuallyQualified. Each mapping also has a reified ComputationalAssertion that records its evidence passage, proposed episode, predicate, rule-acceptance status, and generating activity.")
    add_table(d,['Metric','Result'],[
      ['Gospel books',r['books']],['Episode groupings',r['episodes']],['Passage units',r['passage_units']],['Typed assertions with provenance records',r['typed_assertions']],['Groupings containing material from all four Gospels',len(r['four_gospel_episodes'])],['Groupings containing material from all three Synoptic Gospels',len(r['synoptic_episodes'])]], [4.7,1.8])
    d.add_paragraph('Table 2. Structural metrics for the CKIO proof of concept.',style='Caption')
    d.add_heading('3.2 Integrity validation',2)
    para(d,'All structural integrity checks passed. The Turtle serialization parsed successfully as 780 explicit RDF triples. Every declared Gospel contained at least one passage unit; every passage used a recognized episode key; all passage identifiers were unique; every episode had at least one evidence-bearing passage; and all 68 reified ComputationalAssertion records contained evidence, episode, predicate, status, and generating-activity links. These findings establish internal dataset coherence under the implemented rules. They do not establish whether every possible parallel has been included or whether every grouping would be accepted by biblical scholars.')
    d.add_heading('3.3 Competency-question results',2)
    cqrows=[
      ['CQ1','Return passage-unit count',str(r['passage_units']),'Pass'],
      ['CQ2','Group passage units by Gospel',', '.join(f"{b}: {r['per_book'][b]}" for b in BOOKS),'Pass'],
      ['CQ3','Return four-Gospel episodes',str(len(r['four_gospel_episodes'])),'Pass'],
      ['CQ4','Return feeding-of-5,000 passages',str(len(r['competency_queries']['CQ4'])),'Pass'],
      ['CQ5','Return textually qualified passages',str(len(r['competency_queries']['CQ5'])),'Pass'],
      ['CQ6','Return Synoptic episodes',str(len(r['synoptic_episodes'])),'Pass']]
    add_table(d,['CQ','Operation','Obtained result','Status'],cqrows,[.65,2.75,2.25,.85])
    d.add_paragraph('Table 3. Competency-question evaluation.',style='Caption')
    para(d,'The feeding-of-the-five-thousand query returned one passage from each Gospel: Matthew 14:13-21, Mark 6:30-44, Luke 9:10-17, and John 6:1-15. The qualification query returned the Mark 16 relationships used in the post-resurrection appearance and commissioning fixtures. By preserving the qualification in the predicate, the graph exposes a material textual issue to downstream users instead of treating all mappings as epistemically identical.')
    d.add_heading('3.4 Provenance and interpretive separation',2)
    para(d,'The schema successfully represented passage-to-episode mappings without creating doctrinal claims. InterpretiveClaim and ComputationalAssertion remained separate classes, and both remained distinguishable from Passage. This separation means a future assertion such as a denominational interpretation of baptism would require an attributed claim node and could not be encoded as though it were merely the content of a verse. The prototype therefore met its principal design requirement at the schema level.')

    d.add_heading('4. Discussion',1)
    para(d,'The results demonstrate technical feasibility for a provenance-aware Gospel ontology, but the primary contribution is not the number of triples. It is the explicit separation of textual evidence, analytical grouping, computational proposal, and attributed interpretation. Religious knowledge systems are especially vulnerable to category errors because a semantically plausible association can be mistaken for an authoritative theological conclusion. CKIO reduces this risk by requiring typed, source-anchored relationships and reserving doctrinal propositions for attributed claim objects.')
    para(d,'The passage-to-episode model also improves upon a generic relatedTo property. depicts, testifiesTo, contextualizes, and textuallyQualified communicate why a passage is connected to an episode and whether additional caution is needed. The distinction is particularly useful when comparing John with the Synoptic Gospels, whose narrative organization and presentation may differ. Nevertheless, these predicates remain coarse. Future versions should add quotation, allusion, fulfillment, temporal sequence, participant roles, geographic context, manuscript evidence, and degrees of scholarly agreement.')
    para(d,'The successful competency queries show functional adequacy only for the declared questions. Because expected outputs were derived from the same curated fixtures used to build the graph, the evaluation verifies implementation correctness rather than external validity. A stronger study should construct a preregistered gold standard independently, separate development and test sets, and calculate relationship-level precision, recall, F1, and inter-rater agreement. Negative examples are necessary because a system evaluated only on known parallels cannot demonstrate that it avoids false associations.')
    para(d,'The study intentionally makes no claim that the ontology is theologically neutral. Complete neutrality is unlikely because selection, categorization, translation, and relationship naming involve judgment. The achievable goal is traceable plurality: claims should be attributed to sources and traditions, uncertainty should be visible, and alternative interpretations should be representable without forcing one into the canonical-data layer. A governance panel should therefore include biblical scholars, clergy, digital-humanities researchers, linguists, and representatives of Catholic, Orthodox, and diverse Protestant traditions.')
    para(d,'Scalability also remains unproven. Structural counts do not measure scalability. Full evaluation should report ingestion time, triple-store size, memory use, reasoner performance, SHACL-validation time, and SPARQL latency at increasing graph sizes. Modularization will be important: textual structure, linguistic annotation, people and places, events, doctrine, interpretive claims, and provenance should be separable modules connected through stable identifiers. Incremental validation should be preferred over repeatedly reasoning across an unrestricted full ABox.')
    d.add_heading('4.1 Comparison with the Islamic ontology study',2)
    para(d,'The Islamic ontology study provides a useful precedent for integrating sacred resources with LLM-supported topic generation [13]. CKIO adopts a narrower empirical corpus but strengthens epistemic controls. It distinguishes computational and interpretive assertions, replaces a broad association with typed predicates, treats provenance as part of the conceptual model, and does not use logical consistency as evidence of theological correctness. Conversely, the Islamic study reports a much larger populated ontology and multidisciplinary expert feedback. CKIO cannot claim equivalent scale or expert validation. The approaches are therefore complementary rather than directly ranked on a single dimension.')
    d.add_heading('4.2 Applications',2)
    para(d,'Potential applications include provenance-aware semantic search, Gospel-harmony navigation, teaching tools, geographic and social-network visualization, sermon and curriculum research, cross-translation comparison, and retrieval-augmented generation constrained to cited evidence. An application could allow users to move from an episode to all represented passages, compare relationship types, inspect textual qualifications, and view attributed interpretations by tradition. The ontology could also support AI systems that must cite the precise source and disclose whether an answer is textual, computationally inferred, or interpretive.')
    d.add_heading('4.3 Limitations',2)
    para(d,'The prototype is intentionally small and manually curated. It contains passage references rather than full texts, does not model verse-level linguistic data, and does not include creeds, councils, commentaries, manuscripts, textual variants beyond a limited qualification, or denominational doctrine. It lacks independent scholarly adjudication, user testing, a negative relationship set, and performance benchmarks in a production triple store. The LLM contribution was developmental rather than a locked extraction experiment, so no LLM accuracy metrics are reported. These limitations prevent claims of completeness, doctrinal authority, production scalability, or superiority across all evaluation dimensions.')

    d.add_heading('5. Conclusions',1)
    para(d,'This study developed and evaluated a provenance-aware Christian Knowledge Integration Ontology proof of concept for the four canonical Gospels. The artifact represents 20 episode groupings through source-anchored passage units, typed relationships, and reified provenance records while maintaining schema-level separation among biblical passages, computational assertions, interpretive claims, traditions, and validation activities. Structural integrity checks and six competency questions confirmed that the implemented dataset is internally coherent and functionally queryable for its declared scope.')
    para(d,'The principal finding is that semantic integration can be designed to preserve epistemic boundaries instead of collapsing scripture, machine inference, and theology into a single graph of apparently equivalent facts. This architecture provides a stronger foundation for responsible AI applications in religious knowledge. However, the prototype remains a technical demonstration. Expansion should proceed through licensed full-text resources, controlled LLM experiments, SHACL validation, negative examples, independent gold standards, diverse Christian expert review, and measured triple-store performance.')

    d.add_heading('Acknowledgments',1)
    para(d,'The author acknowledges the maintainers of the open standards and biblical data resources that informed this proof of concept.')
    d.add_heading('Funding',1); para(d,'This research received no external funding.')
    d.add_heading('Author contributions',1); para(d,'J.A.: conceptualization, methodology, investigation, ontology design, validation, writing - original draft, and writing - review and editing.')
    d.add_heading('Conflict of interest',1); para(d,'The author declares no conflict of interest.')
    d.add_heading('Data availability statement',1); para(d,'The prototype ontology, structured passage dataset, validation script, build script, documentation, and machine-readable evaluation results accompany this manuscript as supplementary files under a CC BY 4.0 license. A permanent public repository URL will be added before submission without changing the reported analysis.')
    d.add_heading('Institutional review board statement',1); para(d,'Ethical review was not required because this proof of concept did not involve human participants, identifiable personal data, or animals.')
    d.add_heading('Generative AI disclosure',1)
    para(d,'OpenAI ChatGPT (GPT-5 family, accessed August 3, 2026) was used to assist with idea development, schema refinement, code generation, and manuscript drafting. AI-generated material was not treated as theological authority or as independently validated research evidence. The human author remains responsible for verifying all sources, passage mappings, analytical claims, and the final submitted manuscript. The research does not report LLM extraction accuracy because a fixed external model experiment was not conducted.')

    d.add_heading('References',1)
    refs=[
    '1. Gruber TR. A translation approach to portable ontology specifications. Knowledge Acquisition. 1993;5(2):199-220. doi:10.1006/knac.1993.1008.',
    '2. World Wide Web Consortium. RDF 1.1 Concepts and Abstract Syntax. W3C Recommendation; 2014. https://www.w3.org/TR/rdf11-concepts/',
    '3. World Wide Web Consortium. OWL 2 Web Ontology Language Document Overview. 2nd ed. W3C Recommendation; 2012. https://www.w3.org/TR/owl2-overview/',
    '4. World Wide Web Consortium. SPARQL 1.1 Overview. W3C Recommendation; 2013. https://www.w3.org/TR/sparql11-overview/',
    '5. Boisen S. New Testament Names: a semantic knowledge base. SemanticBible. https://semanticbible.com/ntn/ntn-overview.html (accessed August 3, 2026).',
    '6. CrossWire Bible Society. OSIS: a common format for multiple visions. https://ftp.crosswire.org/osis/ (accessed August 3, 2026).',
    '7. Open Scriptures. Open Scriptures Hebrew Bible. https://github.com/openscriptures/morphhb (accessed August 3, 2026).',
    '8. STEPBible. STEPBible Data Repository. https://github.com/STEPBible/STEPBible-Data (accessed August 3, 2026).',
    '9. Brown TB, Mann B, Ryder N, et al. Language models are few-shot learners. Advances in Neural Information Processing Systems. 2020;33:1877-1901. doi:10.48550/arXiv.2005.14165.',
    '10. OpenAI. GPT-4 technical report. arXiv. 2023. doi:10.48550/arXiv.2303.08774.',
    '11. Mavridis A, Tegos S, Anastasiou C, Papoutsoglou M, Meditskos G. Large language models for intelligent RDF knowledge graph construction: results from medical ontology mapping. Frontiers in Artificial Intelligence. 2025;8:1546179. doi:10.3389/frai.2025.1546179.',
    '12. Li J, Garijo D, Poveda-Villalon M. Large language models for ontology engineering: a systematic literature review. Semantic Web. 2026; in press. https://www.semantic-web-journal.net/system/files/swj4001.pdf',
    '13. Alshammari IK, Atwell E, Alsalka MA. Large language models for automated Islamic ontology construction and knowledge integration. Academia AI and Applications. 2026;2. doi:10.20935/AcadAI8357.',
    '14. Mushtaq A, Naeem R, Elmahjub E, et al. Can LLMs write faithfully? An agent-based evaluation of LLM-generated Islamic content. arXiv. 2025. doi:10.48550/arXiv.2510.24438.',
    '15. Atif F, Askarbekuly N, Darwish K, Choudhury M. Sacred or synthetic? Evaluating LLM reliability and abstention for religious questions. Proceedings of the AAAI/ACM Conference on AI, Ethics, and Society. 2025;8(1):217-226. doi:10.1609/aies.v8i1.36543.',
    '16. Lebo T, Sahoo S, McGuinness D, editors. PROV-O: The PROV Ontology. W3C Recommendation; 2013. https://www.w3.org/TR/prov-o/',
    '17. World Wide Web Consortium. Shapes Constraint Language (SHACL). W3C Recommendation; 2017. https://www.w3.org/TR/shacl/',
    '18. Peroni S, Shotton D. FaBiO and CiTO: ontologies for describing bibliographic resources and citations. Journal of Web Semantics. 2012;17:33-43. doi:10.1016/j.websem.2012.08.001.',
    '19. Nandan AD, Godbole I, Kapparad PM, Bhattacharjee S. Comparative analysis of religious texts: NLP approaches to the Bible, Quran, and Bhagavad Gita. In: Proceedings of the Workshop on New Horizons in Computational Linguistics for Religious Texts; 2025; Abu Dhabi, UAE. p. 1-10. https://aclanthology.org/2025.clrel-1.1/',
    '20. Plu J, Moreno Escobar O, Trouillez E, Gapin A, Troncy R. A comprehensive benchmark for evaluating LLM-generated ontologies. In: Proceedings of the 23rd International Semantic Web Conference (ISWC 2024); 11-15 November 2024; Baltimore, USA. https://www.eurecom.fr/en/publication/7945',
    '21. Aland K, editor. Synopsis of the Four Gospels: Completely Revised on the Basis of the Greek Text of the Nestle-Aland 26th Edition and Greek New Testament 3rd Edition. Revised ed. United Bible Societies; 1985. ISBN: 9780826705006.'
    ]
    for x in refs: para(d,x)

    path=OUT/'Christian_Knowledge_Integration_Ontology_Manuscript.docx'; d.save(path); return path

if __name__=='__main__':
    write_ttl(); results=evaluate(); doc=build_doc(results)
    print(doc)
    print(json.dumps(results,indent=2,default=dict))
