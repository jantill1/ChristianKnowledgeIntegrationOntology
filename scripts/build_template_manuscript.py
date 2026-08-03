from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
TEMPLATE = ROOT / "upload" / "Academia Word Submission Template.docx"
SOURCE = ROOT / "output" / "Christian_Knowledge_Integration_Ontology_Manuscript.docx"
OUTPUT = ROOT / "output" / "Christian_Knowledge_Integration_Ontology_Submission.docx"


def clear_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_text(doc, text, style, bold_prefix=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def main():
    src = Document(SOURCE)
    out = Document(TEMPLATE)
    clear_body(out)

    add_text(out, "Research Article", "Layout Single Article Type")
    add_text(out, src.paragraphs[0].text, "Layout Single Title")
    add_text(out, "John Antill", "Layout Single Author Names")
    add_text(out, "Independent Researcher", "Layout Single Author Affiliations")
    add_text(
        out,
        "* Correspondence: john@littlecrittersnursery.com",
        "Layout Single Corresponding Author",
    )

    abstract_idx = next(i for i, p in enumerate(src.paragraphs) if p.text == "Abstract")
    keywords_idx = next(i for i, p in enumerate(src.paragraphs) if p.text.startswith("Keywords:"))
    add_text(out, "Abstract", "Layout Single Abstract Title")
    abstract_paragraphs = src.paragraphs[abstract_idx + 1:keywords_idx]
    for idx, p in enumerate(abstract_paragraphs):
        style = "Layout Single Abstract Final" if idx == len(abstract_paragraphs) - 1 else "Layout Single Abstract Text"
        prefix = p.text.split(":", 1)[0] + ":" if ":" in p.text else None
        add_text(out, p.text, style, bold_prefix=prefix)
    add_text(out, src.paragraphs[keywords_idx].text, "Layout Single Keywords", bold_prefix="Keywords:")

    table_index = 0
    for p in src.paragraphs[keywords_idx + 1:]:
        text = p.text
        if not text:
            continue
        if p.style.name == "Heading 1":
            add_text(out, text, "Layout Single Heading 1")
        elif p.style.name == "Heading 2":
            add_text(out, text, "Layout Single Heading 2")
        elif p.style.name == "List Number" or p.style.name == "List Bullet":
            add_text(out, text, "Layout Single List Item")
        elif p.style.name == "Caption":
            add_text(out, text, "Layout Single Table Caption")
            tbl = deepcopy(src.tables[table_index]._element)
            out._element.body.insert(len(out._element.body) - 1, tbl)
            table_index += 1
        else:
            add_text(out, text, "Layout Single Paragraph")

    # Apply the template's reference style to every numbered item following References.
    seen_refs = False
    for p in out.paragraphs:
        if p.text == "References":
            seen_refs = True
            continue
        if seen_refs and p.text:
            p.style = out.styles["MDPI_7.1_References"]

    # Keep formal front matter compact and prevent orphaned headings.
    for p in out.paragraphs:
        if p.style.name.startswith("Layout Single Heading"):
            p.paragraph_format.keep_with_next = True
        if p.style.name == "Layout Single Title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = run.font.size or Pt(10)

    out.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
